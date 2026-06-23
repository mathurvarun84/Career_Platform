"""Regression tests for rewriter/parser/validator bugfixes."""

from __future__ import annotations

import json
import tempfile
from pathlib import Path
from unittest.mock import patch

from backend.agents.rewriter import (
    COMPANY_HEADER_START,
    COMPANY_ROLE_START,
    HEADER_END,
    RewriterAgent,
    _ensure_experience_markers,
)
from backend.schemas.common import SectionText, SubEntry
from docx import Document
from parser import _parse_docx
from validator.rewriter_validator import _PLACEHOLDER_RE, _check_placeholder_bleed


def test_parse_docx_handles_missing_paragraph_style() -> None:
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        path = Path(tmp.name)
    doc = Document()
    p = doc.add_paragraph("Bullet without named style")
    p.style = None
    doc.save(str(path))
    try:
        text = _parse_docx(str(path))
        assert "Bullet without named style" in text
    finally:
        path.unlink(missing_ok=True)


def test_parse_docx_skips_duplicate_bullet_normal() -> None:
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        path = Path(tmp.name)
    doc = Document()
    doc.add_paragraph("Owned end-to-end development", style="Normal")
    doc.add_paragraph("Owned end-to-end development", style="List Bullet")
    doc.save(str(path))
    try:
        lines = [ln for ln in _parse_docx(str(path)).splitlines() if ln.strip()]
        owned = [ln for ln in lines if "Owned end-to-end" in ln]
        assert len(owned) == 1
    finally:
        path.unlink(missing_ok=True)


def test_sheroes_experience_markers() -> None:
    text = (
        "July2018–July2019  Delhi, India\n"
        "• Built a gamification system by integrating features\n"
    )
    out = _ensure_experience_markers(
        text, "Sheroes — Software Engineer (July2018–July2019)"
    )
    header = out.split(HEADER_END)[0]
    body = out.split(HEADER_END, 1)[1]
    company_part = header.split(COMPANY_HEADER_START)[1].split(COMPANY_ROLE_START)[0]
    role_part = header.split(COMPANY_ROLE_START)[1]
    assert "Sheroes" in company_part
    assert "July2018" in role_part
    assert "July2018" not in body


def test_strip_unfilled_placeholders() -> None:
    raw = (
        "Owned end-to-end development of [feature/module], "
        "delivering [X%] improvement"
    )
    cleaned = _PLACEHOLDER_RE.sub("", raw).strip()
    assert "[" not in cleaned
    assert "Owned end-to-end" in cleaned


def test_download_blocked_when_placeholder_in_patched_doc() -> None:
    """
    Placeholder tokens in export text must be detected before download sign-off.
    """
    test_cases = [
        ("• Led [N users] through migration", True),
        ("• Reduced latency by [X%]", True),
        ("• Managed [feature/module] deployment", True),
        ("• Led 12 engineers through migration", False),
        ("• Reduced latency by 40%", False),
    ]
    for text, should_match in test_cases:
        has_placeholder = bool(_PLACEHOLDER_RE.search(text))
        assert has_placeholder == should_match, (
            f"{'Should' if should_match else 'Should NOT'} detect placeholder in: {text!r}"
        )


def test_infosys_client_walmart_company_marker_is_infosys() -> None:
    """
    Service company resumes often have 'Client: Walmart' in the experience block.
    The ##COMPANY## marker must always be the employer (Infosys), not the client.
    """
    verbatim = (
        "Infosys — Senior Software Engineer\n"
        "Jan 2019 – Mar 2022\n"
        "Client: Walmart Inc.\n"
        "• Designed inventory allocation system handling 15k QPS\n"
        "• Reduced batch processing time by 40%"
    )
    result = _ensure_experience_markers(
        verbatim, "Infosys — Senior Software Engineer"
    )
    company_part = result.split(COMPANY_HEADER_START)[1].split(COMPANY_ROLE_START)[0]
    assert "Infosys" in company_part, f"Expected Infosys as company, got: {company_part!r}"
    assert "Walmart" not in company_part, f"Client leaked into company field: {company_part!r}"


def test_placeholder_bleed_strips_from_all_styles() -> None:
    variants = {
        "balanced": "Solid rewrite with metrics.",
        "aggressive": "Improved [feature/module] by [X%] and [N users].",
        "top_1_percent": "Also [Xms] latency [INR X Cr] revenue.",
    }
    repaired, anomalies = _check_placeholder_bleed("summary", variants)
    assert repaired["balanced"] == variants["balanced"]
    assert "[" not in repaired["aggressive"]
    assert "[" not in repaired["top_1_percent"]
    assert any("stripped" in a for a in anomalies)


def test_rewrite_monolithic_returns_after_first_success() -> None:
    """Regression: trace call_label must not drop parse/return (double LLM call)."""
    agent = RewriterAgent()
    payload = {
        "balanced": "Balanced summary.",
        "aggressive": "Aggressive summary.",
        "top_1_percent": "Top summary.",
    }
    with patch.object(
        RewriterAgent,
        "_call_llm",
        return_value=json.dumps(payload),
    ) as mock_llm:
        result = agent._rewrite_monolithic(
            "summary",
            "Original summary text.",
            {"rewrite_instruction": "Strengthen.", "missing_keywords": ["Python"]},
        )
    mock_llm.assert_called_once()
    assert result == payload


def test_experience_never_monolithic_without_sub_entries() -> None:
    agent = RewriterAgent()
    gap = {
        "section": "experience",
        "needs_change": True,
        "original_content": f"{COMPANY_HEADER_START}Acme##ROLE##Eng{HEADER_END}\n- bullet",
        "sub_changes": None,
    }
    result, patches = agent._rewrite_section_from_gap("experience", gap, {})
    assert patches == []
    assert "Acme" in result["balanced"]
    assert result["balanced"] == result["aggressive"] == result["top_1_percent"]


def test_experience_with_sub_entries_uses_per_entry_path() -> None:
    entries = [
        SubEntry(label=f"Corp{i}", verbatim_text=f"Corp{i}\nRole\n- b{i}")
        for i in range(4)
    ]
    section = SectionText(
        header="experience",
        full_text="\n\n".join(e.verbatim_text for e in entries),
        sub_entries=entries,
    )
    agent = RewriterAgent()
    gap = {"section": "experience", "needs_change": True, "sub_changes": []}
    result, _ = agent._rewrite_section_from_gap(
        "experience", gap, {"experience": section.model_dump()}
    )
    assert result["balanced"].count(COMPANY_HEADER_START) == 4
