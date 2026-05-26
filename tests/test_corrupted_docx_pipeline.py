"""
Regression: corrupted Varun-style DOCX (split bullets, missing date lines).

Asserts parser + ResumeUnderstandingValidator recover 7 experience entries
and 5 certification entries in ground-truth order.
"""

from __future__ import annotations

import os
from pathlib import Path

import pytest
from docx import Document

from parser import _parse_docx
from validator.resume_understanding_validator import (
    ResumeUnderstandingValidator,
    _detect_sub_entries,
    _extract_all_sections_from_text,
)

# Override via env when running against the real production resume file.
CORRUPTED_RESUME_DOCX = os.environ.get(
    "CORRUPTED_RESUME_DOCX",
    "",
)

EXPERIENCE_COMPANIES_ORDERED = (
    "Flipkart",
    "SmartVizX",
    "Apttus",
    "ClearTax",
    "British Telecom",
    "Microsoft",
    "Mindtree",
)

CERT_EXPECTATIONS = (
    ("IISc", "Agentic AI"),
    ("Google", "Oct 2024"),
    ("University of Queensland", "Aug 2020"),
    ("University at Buffalo", "May 2023"),
    ("Udemy", "May 2020"),
)


def _company_in_label(company: str, label: str) -> bool:
    """Match company tokens after parser camelCase spacing (e.g. SmartVizX → Smart Viz X)."""
    label_norm = label.lower().replace(" ", "")
    aliases = {company.lower().replace(" ", "")}
    if company.lower() == "cleartax":
        aliases.add("cleartax")
    if company.lower() == "british telecom":
        aliases.update({"britishtelecom", "bt"})
    return any(alias in label_norm for alias in aliases)


def _add_paragraph(doc: Document, text: str, style: str) -> None:
    p = doc.add_paragraph(text)
    p.style = style


def _build_corrupted_varun_docx(path: Path) -> None:
    """
    Build a DOCX that mirrors structural corruption:

    - List Bullet vs Normal paragraphs (split mid-sentence continuations)
    - Five experience entries without standalone date-range lines
    - Cert lines using ``Issuer | Mon YYYY`` without ``Certificate`` keywords
    """
    doc = Document()

    _add_paragraph(doc, "Varun Mathur", "Normal")
    _add_paragraph(doc, "varun@example.com | +91-9999999999", "Normal")
    _add_paragraph(doc, "", "Normal")
    _add_paragraph(doc, "E X P E R I E N C E", "Normal")

    # 1. Flipkart — role header, no separate date line; split bullet
    _add_paragraph(
        doc,
        "Engineering Manager  |  Flipkart  —  Bengaluru, KA",
        "Normal",
    )
    _add_paragraph(
        doc,
        "• Lead 5 teams with 32 engineers; previously scaled to 7 teams / ~50 engineers",
        "List Bullet",
    )
    _add_paragraph(
        doc,
        "overhead for two teams and improving overall engineering throughput.",
        "Normal",
    )
    _add_paragraph(doc, "", "Normal")

    # 2. SmartVizX — no date line
    _add_paragraph(
        doc,
        "Head of Engineering  |  SmartVizX  —  Bengaluru, KA",
        "Normal",
    )
    _add_paragraph(
        doc,
        "• Led cross-functional teams for a 3D design SaaS platform.",
        "List Bullet",
    )
    _add_paragraph(doc, "", "Normal")

    # 3. Apttus — inline date on header (one of two entries with a date anchor)
    _add_paragraph(
        doc,
        "Engineering Manager  |  Apttus (via Altran — Consulting Engagement)  —  "
        "Bengaluru, KA  Nov 2018 – Dec 2019",
        "Normal",
    )
    _add_paragraph(
        doc,
        "• Led architecture and delivery of a multi-tenant B2B chatbot.",
        "List Bullet",
    )
    _add_paragraph(doc, "", "Normal")

    # 4. ClearTax — inline date
    _add_paragraph(
        doc,
        "Engineering Manager  |  ClearTax (via Altran — Consulting Engagement)  —  "
        "Bengaluru, KA  Dec 2016 – Oct 2018",
        "Normal",
    )
    _add_paragraph(
        doc,
        "• Conceptualised and delivered the one-page e-filing product;",
        "List Bullet",
    )
    _add_paragraph(
        doc,
        "growing the user base by 50% within one fiscal year.",
        "Normal",
    )
    _add_paragraph(doc, "", "Normal")

    # 5. British Telecom — role header only, no date line
    _add_paragraph(
        doc,
        "Senior Consultant  |  British Telecom  —  Bengaluru, KA",
        "Normal",
    )
    _add_paragraph(
        doc,
        "• Led BOSS application migration to BT Cloud, saving £3.5 million annually.",
        "List Bullet",
    )
    _add_paragraph(
        doc,
        "performance bottlenecks under high-traffic production conditions.",
        "Normal",
    )
    _add_paragraph(doc, "", "Normal")

    # 6. Microsoft — role header only
    _add_paragraph(
        doc,
        "Tech Consultant  |  Microsoft  —  Bengaluru, KA",
        "Normal",
    )
    _add_paragraph(
        doc,
        "• Provided technical leadership for Fortune 500 clients",
        "List Bullet",
    )
    _add_paragraph(
        doc,
        "including Citigroup, Wells Fargo, and VISA — covering financial platforms.",
        "Normal",
    )
    _add_paragraph(doc, "", "Normal")

    # 7. Mindtree — role header only
    _add_paragraph(
        doc,
        "Lead Software Engineer  |  Mindtree  —  Bengaluru, KA",
        "Normal",
    )
    _add_paragraph(
        doc,
        "• Supervised a team of 6 across enterprise clients including Volvo",
        "List Bullet",
    )
    _add_paragraph(
        doc,
        "and Arcelor Mittal; secured a $1M follow-on maintenance contract.",
        "Normal",
    )
    _add_paragraph(doc, "", "Normal")

    _add_paragraph(doc, "C E R T I F I C A T I O N S", "Normal")
    _add_paragraph(
        doc,
        "IISc (Indian Institute of Science) | Agentic AI for Engineers and Managers | Sep 2024",
        "Normal",
    )
    _add_paragraph(doc, "Google | Oct 2024", "Normal")
    _add_paragraph(doc, "University of Queensland | Aug 2020", "Normal")
    _add_paragraph(doc, "University at Buffalo | May 2023", "Normal")
    _add_paragraph(doc, "Udemy | May 2020", "Normal")

    doc.save(str(path))


@pytest.fixture
def corrupted_docx_path(tmp_path: Path) -> Path:
    if CORRUPTED_RESUME_DOCX and Path(CORRUPTED_RESUME_DOCX).is_file():
        return Path(CORRUPTED_RESUME_DOCX)
    path = tmp_path / "varun_corrupted.docx"
    _build_corrupted_varun_docx(path)
    return path


def test_corrupted_docx_parser_and_validator(
    corrupted_docx_path: Path,
) -> None:
    resume_text = _parse_docx(str(corrupted_docx_path))

    # Continuations must be rejoined onto bullets (not standalone lines).
    assert "overhead for two teams" in resume_text
    assert "\noverhead for two teams" not in resume_text
    assert "• Lead 5 teams" in resume_text
    assert resume_text.count("• ") >= 7

    sections = _extract_all_sections_from_text(resume_text)
    exp_blocks = _detect_sub_entries(sections.get("experience", ""), "experience")
    exp_labels = [b["label"] for b in exp_blocks]

    assert len(exp_blocks) == 7, (
        f"expected 7 experience blocks, got {len(exp_blocks)}: {exp_labels}"
    )
    for company in EXPERIENCE_COMPANIES_ORDERED:
        assert any(_company_in_label(company, lbl) for lbl in exp_labels), (
            f"missing company '{company}' in {exp_labels}"
        )
    for idx, company in enumerate(EXPERIENCE_COMPANIES_ORDERED):
        assert _company_in_label(company, exp_labels[idx]), (
            f"order mismatch at index {idx}: expected '{company}' in "
            f"'{exp_labels[idx]}', full order: {exp_labels}"
        )

    # Bullet fragments must not become entry labels.
    bad_labels = ("fiscal year.", "production condit", "overhead for two", "including Citigroup")
    for bad in bad_labels:
        assert not any(lbl.strip().startswith(bad[:20]) for lbl in exp_labels), (
            f"bullet bleed into label: {exp_labels}"
        )

    cert_blocks = _detect_sub_entries(
        sections.get("certifications", ""), "certifications"
    )
    cert_labels = [b["label"] for b in cert_blocks]
    assert len(cert_blocks) == 5, (
        f"expected 5 cert blocks, got {len(cert_blocks)}: {cert_labels}"
    )
    for needle_a, needle_b in CERT_EXPECTATIONS:
        assert any(
            needle_a.lower() in lbl.lower() and needle_b.lower() in lbl.lower()
            for lbl in cert_labels
        ), f"missing cert ({needle_a!r}, {needle_b!r}) in {cert_labels}"

    partial_a1 = {
        "experience_years": 0,
        "seniority": "senior",
        "tech_stack": [],
        "domains": [],
        "has_metrics": True,
        "has_summary": False,
        "sections_present": ["experience", "certifications"],
        "resume_sections": {
            "experience": {"header": "experience", "full_text": "", "sub_entries": []},
            "certifications": {
                "header": "certifications",
                "full_text": "",
                "sub_entries": [],
            },
        },
    }
    repaired = ResumeUnderstandingValidator().validate_and_fix(
        partial_a1, resume_text
    )

    exp_subs = repaired["resume_sections"]["experience"]["sub_entries"]
    assert len(exp_subs) == 7
    for idx, company in enumerate(EXPERIENCE_COMPANIES_ORDERED):
        assert _company_in_label(company, exp_subs[idx]["label"])

    cert_subs = repaired["resume_sections"]["certifications"]["sub_entries"]
    assert len(cert_subs) == 5
    cert_repaired_labels = [s["label"] for s in cert_subs]
    for needle_a, needle_b in CERT_EXPECTATIONS:
        assert any(
            needle_a.lower() in lbl.lower() and needle_b.lower() in lbl.lower()
            for lbl in cert_repaired_labels
        ), f"missing cert ({needle_a!r}, {needle_b!r}) in {cert_repaired_labels}"
