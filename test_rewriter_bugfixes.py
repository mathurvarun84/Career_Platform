"""Regression tests for rewriter/parser/validator bugfixes."""

from __future__ import annotations

import tempfile
from pathlib import Path

from backend.agents.rewriter import (
    COMPANY_HEADER_START,
    COMPANY_ROLE_START,
    HEADER_END,
    RewriterAgent,
    _ensure_experience_markers,
)
from backend.schemas.common import SectionText, SubEntry
from docx import Document
from parser import _parse_docx
from validator.rewriter_validator import _PLACEHOLDER_RE, _check_placeholder_bleed


def test_parse_docx_skips_duplicate_bullet_normal() -> None:
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        path = Path(tmp.name)
    doc = Document()
    doc.add_paragraph("Owned end-to-end development", style="Normal")
    doc.add_paragraph("Owned end-to-end development", style="List Bullet")
    doc.save(str(path))
    try:
        lines = [ln for ln in _parse_docx(str(path)).splitlines() if ln.strip()]
        owned = [ln for ln in lines if "Owned end-to-end" in ln]
        assert len(owned) == 1
    finally:
        path.unlink(missing_ok=True)


def test_sheroes_experience_markers() -> None:
    text = (
        "July2018–July2019  Delhi, India\n"
        "• Built a gamification system by integrating features\n"
    )
    out = _ensure_experience_markers(
        text, "Sheroes — Software Engineer (July2018–July2019)"
    )
    header = out.split(HEADER_END)[0]
    body = out.split(HEADER_END, 1)[1]
    company_part = header.split(COMPANY_HEADER_START)[1].split(COMPANY_ROLE_START)[0]
    role_part = header.split(COMPANY_ROLE_START)[1]
    assert "Sheroes" in company_part
    assert "July2018" in role_part
    assert "July2018" not in body


def test_strip_unfilled_placeholders() -> None:
    raw = (
        "Owned end-to-end development of [feature/module], "
        "delivering [X%] improvement"
    )
    cleaned = _PLACEHOLDER_RE.sub("", raw).strip()
    assert "[" not in cleaned
    assert "Owned end-to-end" in cleaned


def test_placeholder_bleed_strips_from_all_styles() -> None:
    variants = {
        "balanced": "Solid rewrite with metrics.",
        "aggressive": "Improved [feature/module] by [X%] and [N users].",
        "top_1_percent": "Also [Xms] latency [INR X Cr] revenue.",
    }
    repaired, anomalies = _check_placeholder_bleed("summary", variants)
    assert repaired["balanced"] == variants["balanced"]
    assert "[" not in repaired["aggressive"]
    assert "[" not in repaired["top_1_percent"]
    assert any("stripped" in a for a in anomalies)


def test_experience_never_monolithic_without_sub_entries() -> None:
    agent = RewriterAgent()
    gap = {
        "section": "experience",
        "needs_change": True,
        "original_content": f"{COMPANY_HEADER_START}Acme##ROLE##Eng{HEADER_END}\n- bullet",
        "sub_changes": None,
    }
    result, patches = agent._rewrite_section_from_gap("experience", gap, {})
    assert patches == []
    assert "Acme" in result["balanced"]
    assert result["balanced"] == result["aggressive"] == result["top_1_percent"]


def test_experience_with_sub_entries_uses_per_entry_path() -> None:
    entries = [
        SubEntry(label=f"Corp{i}", verbatim_text=f"Corp{i}\nRole\n- b{i}")
        for i in range(4)
    ]
    section = SectionText(
        header="experience",
        full_text="\n\n".join(e.verbatim_text for e in entries),
        sub_entries=entries,
    )
    agent = RewriterAgent()
    gap = {"section": "experience", "needs_change": True, "sub_changes": []}
    result, _ = agent._rewrite_section_from_gap(
        "experience", gap, {"experience": section.model_dump()}
    )
    assert result["balanced"].count(COMPANY_HEADER_START) == 4
